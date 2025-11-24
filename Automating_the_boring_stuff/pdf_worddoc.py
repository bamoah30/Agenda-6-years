#November 19, 2025
'''Pdf and Word Documents:
Pdf and word documents are binary files that require special libraries to read and write.
A binary file contains data encoded in binary (0s and 1s), structured for machine interpretation.
Modules for this section:
-PyPDF2 - for reading and writing PDF files
-python-docx - for reading and writing Word (.docx) files

'''
# PDF
'''.pdf: portable document format, a file format developed by Adobe to present documents consistently across various devices and platforms.
'''

#November 20, 2025
'''Extracting Text from PDFs:
To extract text from a PDF file using PyPDF2, follow these steps:
1. Install PyPDF2 using pip if you haven't already:
   pip install PyPDF2
2. Use the following code to read and extract text from a PDF file:
    import PyPDF2

    # Open the PDF file in read-binary mode
    with open('example.pdf', 'rb') as pdf_file:
    
    # Create a PDF reader object
    pdf_reader = PyPDF2.PdfReader(pdf_file)

    # Get the number of pages in the PDF
    num_pages = len(pdf_reader.pages)

    # Extract text from each page
    for page_num in range(num_pages):
        page = pdf_reader.pages[page_num]
        text = page.extract_text()
        print(f'Page {page_num + 1}:\n{text}\n')

Example usage:
'''
import PyPDF2
# Open the PDF file in read-binary mode
with open('.\\Automating_the_boring_stuff\\how to write a business plan.pdf', 'rb') as pdf_file:
    # Create a PDF reader object
    pdf_reader = PyPDF2.PdfReader(pdf_file)

    # Get the number of pages in the PDF
    num_pages = len(pdf_reader.pages)

    # Extract text from the third page page
    page = pdf_reader.pages[14] # page numbering starts from 0

    text = page.extract_text()
    print(f'Page {num_pages }:{text}')

#Encrypting a PDF:
import PyPDF2

# Open the original PDF file in read-binary mode
with open('.\\Automating_the_boring_stuff\\original.pdf', 'rb') as pdf_file:
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    pdf_writer = PyPDF2.PdfWriter()

    # Copy all pages to the writer object
    for page_num in range(len(pdf_reader.pages)):
        pdf_writer.add_page(pdf_reader.pages[page_num])

    # Encrypt the PDF with a password
    pdf_writer.encrypt('your_password')

    # Write the encrypted PDF to a new file
    with open('.\\Automating_the_boring_stuff\\encrypted.pdf', 'wb') as encrypted_file:
        pdf_writer.write(encrypted_file)
    #NB: Make sure all the pdf files exist in the specified paths before running the code

    pdf_file.close()
    encrypted_file.close()

#NB: After encrypting, you cann't encrpt the same file again without decrypting it first.


#Decrypting a PDF:
import PyPDF2
# Open the encrypted PDF file in read-binary mode
with open('encrypted.pdf', 'rb') as encrypted_file:
    pdf_reader = PyPDF2.PdfReader(encrypted_file)

    # Check if the PDF is encrypted
    if pdf_reader.is_encrypted:
        # Decrypt the PDF with the password
        pdf_reader.decrypt('your_password')

        # Extract text from the first page
        first_page = pdf_reader.pages[0]
        text = first_page.extract_text()
        print(text)

#November 21, 2025
'''Creating PDF:
PyPDF2 cannot write new PDF content from scratch. It can only manipulate existing PDFs.
Thus it can copy, merge, split, encrpyt, drcrypt, and rotate pages of existing PDFs.
To create a PDF documnet using PYPDF2, you need to start with an existing PDF file as a template, and then add or modify content on top of that template.
The stpes below are what we are going to use to create a new PDF document using PyPDF2:
1. Open one or more existing PDFs (the source PDFs) into PdfFileReader
objects.
2. Create a new PdfFileWriter object.
3. Copy pages from the PdfFileReader objects into the PdfFileWriter object.
4. Finally, use the PdfFileWriter object to write the output PDF.

Example:'''
import PyPDF2

print("Here now")

pdf_writer = PyPDF2.PdfWriter()

# First source PDF
with open('.\\Automating_the_boring_stuff\\Machine Learning.pdf', 'rb') as source1_pdf_file:
    pdf_reader1 = PyPDF2.PdfReader(source1_pdf_file)
    for page_num in range(15, 18):  # pages 16–18 (0-indexed)
        pdf_writer.add_page(pdf_reader1.pages[page_num])

# Second source PDF
with open('.\\Automating_the_boring_stuff\\Machine Learning.pdf', 'rb') as source2_pdf_file:
    pdf_reader2 = PyPDF2.PdfReader(source2_pdf_file)
    for page_num in range(3, 6):  # pages 4–6 (0-indexed)
        pdf_writer.add_page(pdf_reader2.pages[page_num])

# Write all collected pages into one output file
with open('.\\Automating_the_boring_stuff\\Creating.pdf', 'wb') as output_pdf_file:
    pdf_writer.write(output_pdf_file)

# November 23, 2025

#Shebang line also known as  hashbang, is the first line of a script file that indicates the interpreter to be used to execute the script.

'''Handling Word Documents:
With the python-docx module, you can read, write, and modify Word (.docx) files.
To install python-docx, use pip:
pip install python-docx
NB: We pip install python-docx, not docx, but we import docx in our code.

A typical .docx file contains:
-Documnet object: represents the entire document
-Paragraph object: represent individual paragraphs within the document
-Run object: represent a contiguous run of text with the same formatting within a paragraph
-Table object: represent tables within the document
'''

#Reading a Word Document:
import docx

# Open the Word document
doc = docx.Document('.\\Automating_the_boring_stuff\\Artificial Construct - Copy.docx')

print(len(doc.paragraphs))  # Print the number of paragraphs in the document   

print(doc.paragraphs[0].text)  # Print the text of the first paragraph

print(doc.paragraphs[1].text)  # Print the text of the second paragraph

print(len(doc.paragraphs[1].runs))  # Print the number of runs in the second paragraph

print(doc.paragraphs[1].runs[0].text)  # Print the text of the first run in the second paragraph

print(doc.paragraphs[1].runs[1].text)  # Print the text of the second run in the second paragraph


'''Styling Paragraphs and Runs:
For word documents, there are three types of text styles:
1. Paragraph styles: applied to entire paragraphs (e.g., Heading 1, Normal)
2. Character styles: applied to specific runs of text within a paragraph (e.g., Bold, Italic)
3. Linked styles: combination of paragraph and character styles

You can modify the style of pragraphs and runs by setting their style attribute to a string value.
NB: The styles for a particular .docx file is stored in the file.

To see available  paragraph styles, use the code:
    from docx import Document

    doc = Document("the path to your .docx file")

    for style in doc.styles:
        print(style.name)

To view available character styles, use the code:
    from docx import Document

    doc = Document("yourfile.docx")

    for style in doc.styles:
        if style.type == 2:  # 2 = WD_STYLE_TYPE.CHARACTER
            print(style.name)

NB: The style types are characterised by numbers:
1 = Paragraphs
2 = Characters
3 = Tables
4 = Lists


Some do's and don'ts when styling paragraphs and runs:
- Don't  set style value to a spaced string like 'Heading 1'; use 'Heading1' instead.
-End the stlye name for runs with 'Char' to indicate character style, e.g., 'EmphasisChar'.

Example: '''
import docx

# Open the Word document
doc = docx.Document('.\\Automating_the_boring_stuff\\Artificial Construct - Copy.docx')

# Modify the style of the first paragraph to 'Bold' 
doc.paragraphs[0].style = 'Normal'

# Modify the style of the first run in the second paragraph to 'EmphasisChar'
doc.paragraphs[1].runs[0].style = 'Default Paragraph Font'

# Save the modified document
doc.save('.\\Automating_the_boring_stuff\\Styled_Document.docx')

# Adding a customized Character style:
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx import Document

doc = Document(".\\Automating_the_boring_stuff\\Styled_Document.docx")

styles = doc.styles
new_style = styles.add_style('MyEmphasis', WD_STYLE_TYPE.CHARACTER) #NB: To create a Customize paragraph,use the same code but change the type to Paragraph

new_style.font.italic = True #type: ignore
new_style.font.size = Pt(12) #type: ignore

doc.paragraphs[1].runs[0].style = 'MyEmphasis'
doc.save("output.docx")

#November 24, 2025

# Creating Word Documents
import docx

doc = docx.Document()
doc.add_paragraph('Hello world!')
paraObj1 = doc.add_paragraph("This is a second paragraph.") #Adding  the first paragraph
paraObj2 = doc.add_paragraph('This is a yet another paragraph.') # Second Paragraph
paraObj1.add_run('This text is being added to the second paragraph')
doc.save('helloworld.docx')


#Adding heading 
doc = docx.Document()
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)

doc.add_heading('Header 3', 3)

doc.add_heading('Header 4', 4)

doc.save('headings.docx')
